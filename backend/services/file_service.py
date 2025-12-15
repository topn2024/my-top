"""
文件处理服务
负责文件上传、验证和文本提取
"""
import os
import logging
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from logger_config import setup_logger, log_service_call
from werkzeug.utils import secure_filename
from typing import Optional, Tuple

logger = setup_logger(__name__)

# 条件导入文档处理库
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available, .doc/.docx support disabled")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not available, PDF support disabled")


class FileService:
    """文件处理服务类"""

    def __init__(self, config):
        """
        初始化文件服务

        Args:
            config: 配置对象，包含UPLOAD_FOLDER和ALLOWED_EXTENSIONS
        """
        self.upload_folder = config.UPLOAD_FOLDER
        self.allowed_extensions = config.ALLOWED_EXTENSIONS
        self.max_file_size = config.MAX_FILE_SIZE

    def allowed_file(self, filename: str) -> bool:
        """
        检查文件扩展名是否允许

        Args:
            filename: 文件名

        Returns:
            bool: 是否允许该文件类型
        """
        if not filename or '.' not in filename:
            return False

        parts = filename.rsplit('.', 1)
        if len(parts) < 2:
            return False

        return parts[1].lower() in self.allowed_extensions

    def validate_file(self, file) -> Tuple[bool, str]:
        """
        验证文件

        Args:
            file: 上传的文件对象

        Returns:
            tuple: (是否有效, 错误消息)
        """
        if not file:
            return False, '没有选择文件'

        if file.filename == '':
            return False, '文件名为空'

        if not self.allowed_file(file.filename):
            return False, f'不支持的文件类型。允许的类型: {", ".join(self.allowed_extensions)}'

        return True, ''


    @log_service_call("保存上传文件")
    def save_file(self, file) -> Tuple[bool, str, Optional[str]]:
        """
        保存上传的文件

        Args:
            file: 上传的文件对象

        Returns:
            tuple: (是否成功, 消息, 文件路径)
        """
        try:
            # 验证文件
            is_valid, error_msg = self.validate_file(file)
            if not is_valid:
                return False, error_msg, None

            # 安全的文件名 - 保留扩展名
            original_filename = file.filename
            logger.info(f"Processing file upload: original filename='{original_filename}'")

            # 提取扩展名
            name, ext = os.path.splitext(original_filename)
            logger.debug(f"Extracted: name='{name}', extension='{ext}'")

            # 安全化文件名（不含扩展名）
            safe_name = secure_filename(name)
            logger.debug(f"Sanitized name: '{safe_name}'")

            # 如果安全化后文件名为空，使用时间戳
            if not safe_name:
                import time
                safe_name = f"upload_{int(time.time())}"
                logger.warning(f"Filename sanitized to empty, using timestamp: '{safe_name}'")

            # 重新组合文件名和扩展名
            filename = safe_name + ext.lower()
            filepath = os.path.join(self.upload_folder, filename)
            logger.info(f"Final filename: '{filename}', saving to: '{filepath}'")

            # 保存文件
            file.save(filepath)
            logger.info(f"✓ File saved successfully: original='{original_filename}' → saved='{filename}'")

            return True, '文件上传成功', filepath

        except Exception as e:
            logger.error(f"Error saving file: {e}", exc_info=True)
            return False, f'文件保存失败: {str(e)}', None


    @log_service_call("提取文件文本")
    def extract_text(self, filepath: str) -> Optional[str]:
        """
        从文件中提取文本

        Args:
            filepath: 文件路径

        Returns:
            提取的文本内容，失败返回None
        """
        logger.info(f"Attempting to extract text from: {filepath}")

        if not os.path.exists(filepath):
            logger.error(f"✗ File not found: {filepath}")
            return None

        # 安全获取文件扩展名
        parts = filepath.rsplit('.', 1)
        if len(parts) < 2:
            logger.error(f"✗ File has no extension: {filepath}")
            return None
        ext = parts[1].lower()
        logger.debug(f"Detected file extension: '{ext}'")

        try:
            text = None
            if ext in ['txt', 'md']:
                logger.info(f"Extracting text file ({ext}): {filepath}")
                text = self._extract_text_file(filepath)
            elif ext == 'pdf':
                logger.info(f"Extracting PDF: {filepath}")
                text = self._extract_pdf(filepath)
            elif ext in ['doc', 'docx']:
                logger.info(f"Extracting DOCX: {filepath}")
                text = self._extract_docx(filepath)
            else:
                logger.warning(f"✗ Unsupported file type: {ext} for file: {filepath}")
                return None

            if text:
                text_preview = text[:100] + '...' if len(text) > 100 else text
                logger.info(f"✓ Successfully extracted {len(text)} characters. Preview: {text_preview}")
            else:
                logger.warning(f"✗ No text extracted from {filepath}")

            return text

        except Exception as e:
            logger.error(f"✗ Error extracting text from {filepath}: {e}", exc_info=True)
            return None

    def _extract_text_file(self, filepath: str) -> Optional[str]:
        """从文本文件提取内容"""
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin1']
        for encoding in encodings:
            try:
                logger.debug(f"Trying encoding: {encoding}")
                with open(filepath, 'r', encoding=encoding) as f:
                    content = f.read()
                    logger.info(f"✓ Successfully decoded with {encoding}, length: {len(content)}")
                    return content
            except UnicodeDecodeError as e:
                logger.debug(f"Encoding {encoding} failed: {e}")
                continue

        logger.error(f"✗ Failed to decode text file with all encodings: {filepath}")
        return None

    def _extract_pdf(self, filepath: str) -> Optional[str]:
        """从PDF文件提取文本"""
        if not PDF_AVAILABLE:
            logger.error("✗ PyPDF2 not available, cannot extract PDF")
            return None

        text = ''
        with open(filepath, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            page_count = len(pdf_reader.pages)
            logger.info(f"PDF has {page_count} pages")

            for i, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                text += page_text + '\n'
                logger.debug(f"Extracted page {i}/{page_count}, length: {len(page_text)}")

        result = text.strip() if text.strip() else None
        if result:
            logger.info(f"✓ PDF extraction successful, total length: {len(result)}")
        else:
            logger.warning(f"✗ PDF extraction returned empty text")
        return result

    def _extract_docx(self, filepath: str) -> Optional[str]:
        """从DOCX文件提取文本"""
        if not DOCX_AVAILABLE:
            logger.error("✗ python-docx not available, cannot extract DOCX")
            return None

        doc = Document(filepath)
        paragraph_count = len(doc.paragraphs)
        logger.info(f"DOCX has {paragraph_count} paragraphs")

        text = '\n'.join([para.text for para in doc.paragraphs])
        result = text.strip() if text.strip() else None

        if result:
            logger.info(f"✓ DOCX extraction successful, total length: {len(result)}")
        else:
            logger.warning(f"✗ DOCX extraction returned empty text")
        return result

    def delete_file(self, filepath: str) -> bool:
        """
        删除文件

        Args:
            filepath: 文件路径

        Returns:
            bool: 是否成功删除
        """
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
                logger.info(f"File deleted: {filepath}")
                return True
            else:
                logger.warning(f"File not found for deletion: {filepath}")
                return False
        except Exception as e:
            logger.error(f"Error deleting file {filepath}: {e}", exc_info=True)
            return False
